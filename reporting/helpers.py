import base64
import copy
import json
import re
import shutil
import subprocess
import textwrap
import uuid
from collections import Counter, OrderedDict, defaultdict
from itertools import groupby, zip_longest
from pathlib import Path
from statistics import mean

import plotly.colors as pc
import plotly.graph_objects as go
import plotly.io as pio
import psutil
from django.db.models import Avg, Count, F, Min, OuterRef, Q, Subquery
from django.db.models.functions import Floor
from django.forms.models import model_to_dict
from django.utils.translation import gettext as _
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from parler.models import TranslationDoesNotExist

from governanceplatform.helpers import is_user_regulator
from securityobjectives.models import (
    Domain,
    MaturityLevel,
    SecurityObjective,
    SecurityObjectivesInStandard,
    SecurityObjectiveStatus,
    StandardAnswer,
)

from .globals import TRANSLATIONS_CONTEXT
from .models import (
    AssetData,
    LogReporting,
    RecommendationData,
    RiskData,
    ServiceStat,
    ThreatData,
    VulnerabilityData,
)


def get_so_data(cleaned_data):
    def get_latest_answers(company, sector, year):
        latest_submit_date = (
            StandardAnswer.objects.filter(
                standard__id=standard_id,
                submitter_company=OuterRef("submitter_company"),
                sectors__id__in=[sector["id"]],
                year_of_submission=year,
                status="PASSM",
            )
            .order_by("-submit_date")
            .values("submit_date")[:1]
        )

        queryset = StandardAnswer.objects.filter(
            standard__id=standard_id,
            sectors__id__in=[sector["id"]],
            year_of_submission=year,
            status="PASSM",
            submit_date=Subquery(latest_submit_date),
        ).distinct()

        if company is not None:
            queryset = queryset.filter(submitter_company__id=company["id"])

        return queryset

    def calculate_evolution(previous_score, current_score):
        if previous_score is not None:
            previous_score = round(previous_score, 1)
            current_score = round(current_score, 1)
            if previous_score > current_score:
                return False
            if previous_score < current_score:
                return True
            return "="
        return None

    def build_dict_scores(
        queryset,
        *,
        target_dict,
        key_func,
        keys_queryset=None,
        sector_values=False,
        ndigits=0,
    ):
        score_field = "score_value"
        if not queryset and keys_queryset:
            for key in keys_queryset:
                entry = target_dict.setdefault(
                    str(key),
                    {
                        "label": str(key),
                        "score": [],
                        "evolution": [],
                        "sector_avg": [],
                    },
                )
                if sector_values:
                    entry["sector_avg"].append(None)
                else:
                    entry["score"].append({"value": None, "color": "#FFFFFF"})
                    entry["evolution"].append(None)

            return

        for item in queryset.iterator():
            key = str(key_func(item))
            entry = target_dict.setdefault(
                key,
                {
                    "label": key,
                    "score": [],
                    "evolution": [],
                    "sector_avg": [],
                },
            )

            value = item[score_field] if isinstance(item, dict) else getattr(item, score_field)
            score = round(value, ndigits) if value is not None else None

            if sector_values:
                entry["sector_avg"].append(score)
                continue

            previous_score = entry["score"][-1]["value"] if entry["score"] else None
            evolution = calculate_evolution(previous_score, score)
            entry["score"].append(
                {
                    "value": score,
                    "color": (get_gradient_color(score, so_color_palette) if is_last_year else "#FFFFFF"),
                }
            )
            entry["evolution"].append(evolution)

    def build_radar_data(data_dict, sector_avg_field=None):
        radar_data = defaultdict(list)
        score_field = "score"
        for index, year in enumerate(years):
            for __key, values in data_dict.items():
                year_label = f"{year}"
                if index < len(values[score_field]):
                    radar_data[year_label].append(values[score_field][index]["value"])

                if sector_avg_field and is_last_year:
                    sector_label = f"{sector_avg_translation} {reference_year}"
                    if index < len(values[score_field]):
                        radar_data[sector_label].append(values[sector_avg_field][index])

        return dict(radar_data)

    def build_bar_chart_by_level_data(counts, label):
        levels_count = maturity_levels_queryset.count()
        score_list = [counts.get(level, 0) for level in range(levels_count)]
        bar_chart_data_by_level[label] = score_list

    def get_grouped_scores(order="asc") -> dict:
        scores = sector_scores_queryset.order_by("score_value", "min_position")
        so_ids = scores.values_list("security_objective", flat=True)
        so_map = {so.pk: str(so) for so in SecurityObjective.objects.filter(pk__in=so_ids)}
        company_scores_map = (
            {fs.security_objective_id: fs.score_value for fs in floored_company_queryset.filter(security_objective__in=so_ids)}
            if floored_company_queryset
            else {}
        )

        grouped_scores = {
            str(round(score_value, 1)): [
                {
                    "name": so_map[s["security_objective"]],
                    "score_sector": round(score_value, 1),
                    "score_company": round(company_scores_map.get(s["security_objective"], 0), 1),
                }
                for s in group
            ]
            for score_value, group in groupby(scores, key=lambda s: s["score_value"])
        }

        reverse = order == "desc"
        sorted_keys = sorted(grouped_scores.keys(), key=float, reverse=reverse)
        sliced_keys = sorted_keys[:top_ranking]

        return {key: grouped_scores[key] for key in sliced_keys}

    company = cleaned_data["company"]
    sector = cleaned_data["sector"]
    reference_year = cleaned_data["reference_year"]
    years = cleaned_data["years"]
    top_ranking = cleaned_data["top_ranking"]
    standard_id = cleaned_data["standard_id"]
    maturity_levels_queryset = MaturityLevel.objects.filter(standard_id=standard_id).order_by("level")
    domains_queryset = Domain.objects.filter(standard__id=standard_id)
    security_objective_queryset = SecurityObjectivesInStandard.objects.filter(standard__id=standard_id)
    so_color_palette = list(maturity_levels_queryset.values_list("level", "color"))
    maturity_levels = [{"level": ml.level, "label": str(ml), "color": ml.color} for ml in maturity_levels_queryset]
    maturity_levels_labels = [str(ml) for ml in maturity_levels_queryset]
    sector_so_by_year_desc = OrderedDict()
    sector_so_by_year_asc = OrderedDict()
    bar_chart_data_by_level = defaultdict()
    company_so_by_year = defaultdict(dict)
    company_so_by_domain = defaultdict(dict)
    company_so_by_priority = defaultdict(dict)
    radar_chart_data_by_domain = defaultdict()
    radar_chart_data_by_year = defaultdict()
    sector_avg_translation = TRANSLATIONS_CONTEXT["sector_average"]

    for year in years:
        last_answers = get_latest_answers(company, sector, year)

        if not last_answers:
            continue

        security_objective_by_priority_queryset = (
            SecurityObjectiveStatus.objects.filter(standard_answer__in=last_answers)
            .annotate(
                score_value=Floor(F("score")),
                min_priority=Min("security_objective__standard_link__priority"),
            )
            .order_by("score_value", "min_priority")
        )

        break

    for year in years:
        is_last_year = year == reference_year
        company_so_by_level = defaultdict(list)
        latest_answers = get_latest_answers(company, sector, year)

        # Company Querysets
        floored_company_queryset = SecurityObjectiveStatus.objects.filter(standard_answer__in=latest_answers).annotate(
            score_value=Floor(F("score"))
        )

        so_domain_company_queryset = (
            floored_company_queryset.values("security_objective__domain")
            .annotate(score_value=Avg("score"))
            .order_by("security_objective__domain__position")
        )

        sorted_company_queryset = floored_company_queryset.annotate(min_position=Min("security_objective__standard_link__position"))

        company_by_priority_queryset = floored_company_queryset.annotate(
            min_priority=Min("security_objective__standard_link__priority")
        ).order_by("score_value", "min_priority")

        aggregated_scores_queryset = floored_company_queryset.values("score_value").annotate(count=Count("score_value"))

        # Sector Querysets

        latest_answers_sector = get_latest_answers(None, sector, year)

        sector_queryset = SecurityObjectiveStatus.objects.filter(standard_answer__in=latest_answers_sector)

        sector_score_by_domain_queryset = (
            sector_queryset.values("security_objective__domain")
            .annotate(score_value=Avg("score"))
            .order_by("security_objective__domain__position")
        )

        sector_scores_queryset = sector_queryset.values("security_objective").annotate(
            score_value=Floor(Avg("score")),
            min_position=Min("security_objective__standard_link__position"),
        )

        # Dictionaries

        # by_level
        levels_map = {ml.level: f"{ml.level} ({str(ml)})" for ml in maturity_levels_queryset}
        for score_data in sorted_company_queryset.order_by("score_value", "min_position"):
            level = levels_map.get(int(score_data.score_value), "Unknown")
            security_objective = str(score_data.security_objective)
            company_so_by_level[level].append(security_objective)

        headers = list(company_so_by_level.keys())
        columns = list(company_so_by_level.values())

        rows = [{headers[i]: cell or "" for i, cell in enumerate(row)} for row in zip_longest(*columns, fillvalue="")]

        company_so_by_level = {
            "headers": headers,
            "rows": rows,
        }

        # by_domain
        build_dict_scores(
            so_domain_company_queryset,
            target_dict=company_so_by_domain,
            key_func=lambda x: Domain.objects.get(pk=x["security_objective__domain"]),
            keys_queryset=domains_queryset,
            ndigits=1,
        )

        # by_domain [sector]
        build_dict_scores(
            sector_score_by_domain_queryset,
            target_dict=company_so_by_domain,
            key_func=lambda x: Domain.objects.get(pk=x["security_objective__domain"]),
            keys_queryset=domains_queryset,
            sector_values=True,
            ndigits=1,
        )

        # by_year
        build_dict_scores(
            sorted_company_queryset.order_by("min_position"),
            target_dict=company_so_by_year,
            key_func=lambda x: x.security_objective,
            keys_queryset=security_objective_queryset,
        )

        # by_priority
        build_dict_scores(
            company_by_priority_queryset,
            target_dict=company_so_by_priority,
            key_func=lambda x: x.security_objective,
            keys_queryset=security_objective_by_priority_queryset,
        )

        # Bar chart by level
        company_counts = Counter({score_data["score_value"]: score_data["count"] for score_data in aggregated_scores_queryset})

        sector_counts = Counter(
            {
                score_data["score_value"]: sector_scores_queryset.filter(score_value=score_data["score_value"]).count()
                for score_data in sector_scores_queryset
            }
        )

        build_bar_chart_by_level_data(
            counts=company_counts,
            label=f"{year}",
        )

        build_bar_chart_by_level_data(
            counts=sector_counts,
            label=f"{sector_avg_translation} {year}",
        )

        # Sector asc and desc lists by score
        sector_so_by_year_asc[year] = get_grouped_scores(order="asc")
        sector_so_by_year_desc[year] = get_grouped_scores(order="desc")

    radar_chart_data_by_domain_with_sector_avg = build_radar_data(company_so_by_domain, "sector_avg")
    radar_chart_data_by_domain = build_radar_data(company_so_by_domain)
    radar_chart_data_by_year = build_radar_data(company_so_by_year)
    return {
        "domains": [str(domain) for domain in company_so_by_domain.keys()],
        "maturity_levels": maturity_levels,
        "maturity_levels_labels": maturity_levels_labels,
        "unique_codes_list": [so.security_objective.unique_code for so in sorted_company_queryset.order_by("min_position")],
        "max_of_company_count": max(company_counts.values()),
        "bar_chart_data_by_level": dict(sort_legends(bar_chart_data_by_level)),
        "company_so_by_level": company_so_by_level,
        "company_so_by_domain": [item for item in company_so_by_domain.values()],
        "company_so_by_year": [item for item in company_so_by_year.values()],
        "company_so_by_priority": [item for item in company_so_by_priority.values()],
        "sector_so_by_year_desc": dict(sector_so_by_year_desc),
        "sector_so_by_year_asc": dict(sector_so_by_year_asc),
        "radar_chart_data_by_domain": radar_chart_data_by_domain,
        "radar_chart_data_by_domain_with_sector_avg": radar_chart_data_by_domain_with_sector_avg,
        "radar_chart_data_by_year": radar_chart_data_by_year,
    }


def get_risk_data(cleaned_data):
    def build_risk_average_data(service):
        service_stat_queryset = ServiceStat.objects.filter(
            service=service,
            company_reporting__year=year,
            company_reporting__sector__id=sector_id,
        )

        average_risks_by_sector = service_stat_queryset.aggregate(avg_current_risks=Avg("avg_current_risks"))["avg_current_risks"]

        servicestat_by_company = service_stat_queryset.filter(
            company_reporting__company__id=company_id,
        )

        average_risks_by_company = servicestat_by_company.first().avg_current_risks if servicestat_by_company else 0

        for label in labels:
            data_by_risk_average.setdefault(label, [])

        data_by_risk_average[company_label].append(round_value(average_risks_by_company))

        data_by_risk_average[sector_label].append(round_value(average_risks_by_sector))

    def build_high_risk_data(service):
        # Querysets
        risk_data_service_queryset = RiskData.objects.filter(
            service__service=service,
            service__company_reporting__year=year,
            service__company_reporting__sector__id=sector_id,
            max_risk__gt=threshold_for_high_risk,
        ).exclude(risk_treatment="UNTRE")

        service_stat_queryset = ServiceStat.objects.filter(
            service=service,
            company_reporting__year=year,
            company_reporting__company__id=company_id,
            company_reporting__sector__id=sector_id,
        )

        high_risks_by_company = risk_data_service_queryset.filter(
            service__company_reporting__company__id=company_id,
        ).aggregate(count=Count("id"), max_risk_avg=Avg("max_risk"))

        high_risks_average_by_sector = risk_data_service_queryset.aggregate(max_risk_avg=Avg("max_risk"))

        # High risk rate data
        total_high_risks = high_risks_by_company["count"]

        total_risk = service_stat_queryset.first().total_treated_risks if service_stat_queryset else 0

        rate = (total_high_risks / total_risk) if total_risk else 0

        data_by_high_risk_rate.setdefault(year, {"rate_labels": [], "rate_values": []})

        data_by_high_risk_rate[year]["rate_labels"].append(f"{total_high_risks} / {total_risk}")
        data_by_high_risk_rate[year]["rate_values"].append(rate)

        # High risk average data
        for label in labels:
            data_by_high_risk_average.setdefault(label, [])

        data_by_high_risk_average[company_label].append(round_value(high_risks_by_company["max_risk_avg"]))
        data_by_high_risk_average[sector_label].append(round_value(high_risks_average_by_sector["max_risk_avg"]))

    def build_evolution_highest_risks_data(company_reporting):
        risks_data = (
            RiskData.objects.filter(
                service__company_reporting__id=company_reporting["id"],
            )
            .exclude(risk_treatment="UNTRE")
            .annotate(total_impact=F("impact_c") + F("impact_i") + F("impact_a"))
            .order_by(
                "-max_risk",
                "-total_impact",
                "-threat_value",
                "-vulnerability_value",
                "asset__translations__name",
            )
        )

        top_ranking_distinct_risks = get_top_ranking_distinct_risks(risks_data, top_ranking)

        data_evolution_highest_risks[f"{reference_year}"] = [round_value(risk.max_risk) for risk in top_ranking_distinct_risks]

        for year in years:
            for risk in top_ranking_distinct_risks:
                uuid = risk.uuid
                if uuid not in risks_top_ranking:
                    risks_top_ranking_ids.append(f"R{risk.id}")
                    risks_top_ranking[uuid] = model_to_dict(
                        risk,
                        exclude=[
                            "service",
                            "asset",
                            "threat",
                            "vulnerability",
                            "recommendations",
                        ],
                    )
                    risks_top_ranking[uuid]["treatment"] = risk.get_risk_treatment_display()
                    risks_top_ranking[uuid]["service"] = str(risk.service.service)
                    risks_top_ranking[uuid]["asset"] = str(risk.asset)
                    risks_top_ranking[uuid]["threat"] = str(risk.threat)
                    risks_top_ranking[uuid]["vulnerability"] = str(risk.vulnerability)
                    risks_top_ranking[uuid]["risks_values"] = {
                        reference_year: {
                            "c": ({"value": risk.risk_c if risk.risk_c > -1 else None}),
                            "i": ({"value": risk.risk_i if risk.risk_i > -1 else None}),
                            "a": ({"value": risk.risk_a if risk.risk_a > -1 else None}),
                            "max": risk.max_risk or "-",
                        }
                    }

                if year == reference_year:
                    continue

                try:
                    risk = RiskData.objects.get(  # noqa: PLW2901
                        uuid=uuid,
                        service__service=risk.service.service,
                        service__company_reporting__year=year,
                        service__company_reporting__company__id=company_id,
                        service__company_reporting__sector__id=sector_id,
                    )

                    max_risk = risk.max_risk

                    risk_values_dict = {
                        "c": ({"value": risk.risk_c if risk.risk_c > -1 else None}),
                        "i": ({"value": risk.risk_i if risk.risk_i > -1 else None}),
                        "a": ({"value": risk.risk_a if risk.risk_a > -1 else None}),
                        "max": max_risk or "-",
                    }

                except RiskData.DoesNotExist:
                    max_risk = 0
                    risk_values_dict = {
                        "c": None,
                        "i": None,
                        "a": None,
                        "max": "-",
                    }

                risks_top_ranking[uuid]["risks_values"][year] = risk_values_dict

                data_evolution_highest_risks[f"{year}"].append(round_value(max_risk))

    def get_top_by_occurrence(model):
        qs = (
            model.objects.filter(
                riskdata__service__company_reporting__company__id=company_id,
                riskdata__service__company_reporting__year=year,
                riskdata__service__company_reporting__sector__id=sector_id,
            )
            .annotate(
                occurrences=Count(
                    "riskdata",
                    filter=~Q(riskdata__risk_treatment="UNTRE"),
                )
            )
            .order_by("-occurrences")
        )

        grouped = OrderedDict()

        for obj in qs:
            grouped.setdefault(obj.occurrences, []).append(str(obj))

        return sorted(
            grouped.items(),
            key=lambda x: x[0],
            reverse=True,
        )[:top_ranking]

    def build_evolution_recommendations_data():
        recommendations_by_year = RecommendationData.objects.filter(
            riskdata__service__company_reporting__year=year,
            riskdata__service__company_reporting__company__id=company_id,
            riskdata__service__company_reporting__sector__id=sector_id,
        )

        if recommendations_by_year:
            for recommendation in recommendations_by_year:
                color = "#FFFFFF"
                recommendation_key = generate_combined_uuid([recommendation.code, recommendation.description])
                recommendations_evolution.setdefault(
                    recommendation_key,
                    {
                        "code": recommendation.code,
                        "description": recommendation.description,
                        "due_date": recommendation.due_date.strftime("%d/%m/%Y"),
                    },
                )

                previous_year = year - 1
                previous_due_date = recommendations_evolution[recommendation_key].get(previous_year)

                if previous_due_date:
                    status = _("Postponed")
                    color = "#FFC000"
                    if previous_due_date == recommendation.due_date and recommendation.due_date.year == previous_year:
                        status = _("To check")
                else:
                    if is_last_year:
                        status = _("New")
                    else:
                        status = _("Closed")
                        color = "#00B050"
                recommendations_evolution[recommendation_key]["status"] = status
                recommendations_evolution[recommendation_key]["color"] = color
                recommendations_evolution[recommendation_key][year] = recommendation.due_date

    company = cleaned_data["company"]
    company_id = company["id"]
    sector = cleaned_data["sector"]
    sector_id = sector["id"]
    reference_year = cleaned_data["reference_year"]
    years = cleaned_data["years"]
    threshold_for_high_risk = cleaned_data["threshold_for_high_risk"]
    top_ranking = cleaned_data["top_ranking"]
    company_reporting = cleaned_data["company_reporting"]
    data_by_risk_average = defaultdict()
    data_by_high_risk_rate = defaultdict()
    data_by_high_risk_average = defaultdict()
    data_evolution_highest_risks = defaultdict(list)
    risks_top_ranking = OrderedDict()
    risks_top_ranking_ids = []
    risks_stats_by_year = OrderedDict()
    recommendations_evolution = defaultdict(dict)
    services_list = AssetData.objects.filter(servicestat__company_reporting__id=company_reporting["id"]).order_by("id").distinct()
    operator_services = [str(service) for service in services_list]
    operator_services_with_all = [_("All services")] + operator_services

    for year in years:
        is_last_year = year == reference_year
        company_label = f"{year}"
        sector_avg_translation = TRANSLATIONS_CONTEXT["sector_average"]
        sector_label = f"{sector_avg_translation} {year}"
        labels = [company_label, sector_label]

        for service in services_list:
            build_risk_average_data(service)
            build_high_risk_data(service)

        # Score mean for all services
        for label in labels:
            data_by_risk_average[label].insert(0, round_value(mean(data_by_risk_average[label])))
            data_by_high_risk_average[label].insert(0, round_value(mean(data_by_high_risk_average[label])))

        # Stats by year data
        risks_stats_by_year.setdefault(year, {})
        service_stat_by_year_qs = ServiceStat.objects.filter(
            company_reporting__year=year,
            company_reporting__company__id=company_id,
            company_reporting__sector__id=sector_id,
        )
        total_high_risks_qs = (
            RiskData.objects.filter(
                service__company_reporting__company__id=company_id,
                service__company_reporting__year=year,
                service__company_reporting__sector__id=sector_id,
                max_risk__gt=threshold_for_high_risk,
            )
            .exclude(risk_treatment="UNTRE")
            .aggregate(count=Count("id"), max_risk_avg=Avg("max_risk"))
        )
        if service_stat_by_year_qs:
            risks_stats_by_year[year] = model_to_dict(
                service_stat_by_year_qs.first(),
                exclude=["service", "company_reporting"],
            )
            risks_stats_by_year[year]["total_high_risks_treated"] = total_high_risks_qs["count"]
            risks_stats_by_year[year]["avg_high_risks_treated"] = total_high_risks_qs["max_risk_avg"]

            for key in [
                "avg_current_risks",
                "avg_residual_risks",
                "avg_high_risks_treated",
            ]:
                risks_stats_by_year[year][key] = round_value(risks_stats_by_year[year][key])

        if is_last_year:
            build_evolution_highest_risks_data(company_reporting)

            # Top of threats by occurence
            top_threats = get_top_by_occurrence(
                ThreatData,
            )

            # Top of vulnerabilities by occurence
            top_vulnerabilities = get_top_by_occurrence(
                VulnerabilityData,
            )

        build_evolution_recommendations_data()

    return {
        "threshold_for_high_risk": threshold_for_high_risk,
        "data_by_risk_average": dict(sort_legends(data_by_risk_average)),
        "data_by_high_risk_rate": dict(data_by_high_risk_rate),
        "data_by_high_risk_average": dict(sort_legends(data_by_high_risk_average)),
        "data_evolution_highest_risks": dict(sort_legends(data_evolution_highest_risks)),
        "data_risks_top_ranking": list(values for _uuid, values in risks_top_ranking.items()),
        "risks_top_ranking_ids": risks_top_ranking_ids,
        "risks_stats_by_year": dict(risks_stats_by_year),
        "top_threats": dict(top_threats),
        "top_vulnerabilities": dict(top_vulnerabilities),
        "recommendations_evolution": list(reco for _uuid, reco in recommendations_evolution.items()),
        "operator_services": operator_services,
        "operator_services_with_all": operator_services_with_all,
    }


def get_charts(so_data, risk_data, colors):
    charts = {
        "chart_security_objectives_by_level": generate_bar_chart(
            so_data["bar_chart_data_by_level"],
            so_data["maturity_levels_labels"],
            colors,
        ),
        "chart_evolution_security_objectives_by_domain": generate_radar_chart(
            so_data["radar_chart_data_by_domain"],
            so_data["domains"],
            so_data["maturity_levels_labels"],
            colors,
        ),
        "chart_evolution_security_objectives_by_domain_with_sector_avg": generate_radar_chart(
            so_data["radar_chart_data_by_domain_with_sector_avg"],
            so_data["domains"],
            so_data["maturity_levels_labels"],
            colors,
        ),
        "chart_evolution_security_objectives": generate_radar_chart(
            so_data["radar_chart_data_by_year"],
            so_data["unique_codes_list"],
            so_data["maturity_levels_labels"],
            colors,
        ),
        "chart_average_risk_level": generate_bar_chart(
            risk_data["data_by_risk_average"],
            risk_data["operator_services_with_all"],
            colors,
        ),
        "chart_high_risk_rate": generate_bar_chart(
            risk_data["data_by_high_risk_rate"],
            risk_data["operator_services"],
            colors,
            True,
        ),
        "chart_average_high_risk_level": generate_bar_chart(
            risk_data["data_by_high_risk_average"],
            risk_data["operator_services_with_all"],
            colors,
        ),
        "chart_evolution_highest_risks": generate_bar_chart(
            risk_data["data_evolution_highest_risks"],
            risk_data["risks_top_ranking_ids"],
            colors,
        ),
    }

    return {key: convert_graph_to_base64(fig) for key, fig in charts.items()}


def sort_legends(data):
    return sorted(
        data.items(),
        key=lambda x: (x[0] != TRANSLATIONS_CONTEXT["sector_average"], x[0]),
    )


def round_value(value):
    if value is None:
        return 0

    if isinstance(value, int):
        return value

    rounded_value = round(value, 2)

    return int(rounded_value) if rounded_value.is_integer() else float(f"{rounded_value:.2f}")


def get_nested_attr(obj, attr):
    attributes = attr.split(".")
    for attribute in attributes:
        try:
            obj = getattr(obj, attribute)
        except TranslationDoesNotExist:
            obj = str(obj)
    return obj


def get_top_ranking_distinct_risks(data, top_ranking):
    seen_uuids = set()
    distinct = []
    for obj in data:
        array_uuid = [
            obj.service.service.uuid,
            str(obj.asset),
            obj.threat.uuid,
            obj.vulnerability.uuid,
        ]
        combined_uuid = generate_combined_uuid(array_uuid)
        if combined_uuid not in seen_uuids:
            seen_uuids.add(combined_uuid)
            distinct.append(obj)

        if len(distinct) >= top_ranking:
            break
    return distinct


def generate_combined_uuid(array_uuid: list[str]) -> uuid.UUID:
    combined = "".join(str(i_uuid) for i_uuid in array_uuid)
    new_uuid = uuid.uuid3(uuid.NAMESPACE_DNS, combined)

    return str(new_uuid)


def create_entry_log(user, project, action):
    role = user.groups.first().name if user.groups.exists() else ""
    entity_name = ""

    if is_user_regulator(user):
        regulator = user.regulators.first()
        entity_name = regulator.name if regulator else ""

    log = LogReporting.objects.create(
        user=user,
        project=project,
        action=action,
        role=role,
        entity_name=entity_name,
    )
    log.save()


def generate_bar_chart(data, labels, colors, is_rate=False):
    fig = go.Figure()
    labels = text_wrap(labels)
    bar_colors_palette = get_chart_color_palette(colors)
    avg_index = 0
    bar_index = 0
    legend_orientation = "h"
    legend_x = 0.5
    legend_xanchor = "center"

    for name, values in data.items():
        if is_rate:
            rate_labels = values["rate_labels"]
            rate_values = values["rate_values"]

        group_name = str(name)[-4:]
        is_avg = str(_("average")) in str(name).lower()

        if is_avg:
            legend_orientation = "v"
            legend_x = 0
            legend_xanchor = "left"
            marker_color = lighten_color(bar_colors_palette[avg_index])
            fig.add_trace(
                go.Scatter(
                    x=labels,
                    y=rate_values if is_rate else values,
                    name=str(name),
                    mode="markers",
                    marker=dict(size=12, symbol="diamond", color=marker_color),
                    offsetgroup=group_name,
                    legendgroup=group_name,
                    legend="legend2",
                ),
            )
            avg_index += 1

        else:
            marker_color = bar_colors_palette[bar_index]
            fig.add_trace(
                go.Bar(
                    x=labels,
                    y=rate_values if is_rate else values,
                    name=str(name),
                    marker_color=marker_color,
                    text=rate_labels if is_rate else values,
                    textposition="outside",
                    offsetgroup=group_name,
                    legendgroup=group_name,
                    legend="legend",
                ),
            )
            bar_index += 1

    fig.update_layout(
        hovermode="closest",
        barmode="group",
        scattermode="group",
        bargroupgap=0.2,
        xaxis=dict(
            linecolor="black",
        ),
        yaxis=dict(
            showgrid=True,
            gridcolor="lightgray",
            linecolor="black",
            rangemode="nonnegative",
        ),
        plot_bgcolor="rgba(0,0,0,0)",
        showlegend=True,
        legend=dict(
            orientation=legend_orientation,
            x=legend_x,
            y=-0.10,
            xanchor=legend_xanchor,
            yanchor="top",
            traceorder="normal",
            itemwidth=70,
            valign="middle",
        ),
        legend2=dict(
            orientation="v",
            x=0.5,
            y=-0.10,
            xanchor="left",
            yanchor="top",
            traceorder="normal",
            itemwidth=70,
            valign="middle",
        ),
        margin=dict(l=0, r=0, t=0, b=80),
    )

    if is_rate:
        fig.update_layout(yaxis_tickformat=".0%")

    return fig


def generate_radar_chart(data, labels, levels, colors):
    fig = go.Figure()
    labels = text_wrap(labels)
    line_colors_palette = get_chart_color_palette(colors)
    index = 0

    for name, values in data.items():
        if not values[0]:
            continue
        line_style = "solid"
        line_color = line_colors_palette[index]
        marker_color = lighten_color(line_colors_palette[index])
        symbol = "circle"

        if str(_("average")) in str(name).lower():
            line_style = "dash"
            line_color = "#666666"
            marker_color = "#222A2A"
            symbol = "triangle-up"

        r_values = [float(v) for v in values]

        fig.add_trace(
            go.Scatterpolar(
                r=r_values + [r_values[0]],
                theta=labels + [labels[0]],
                name=str(name),
                fillcolor="rgba(0,0,0,0)",
                marker=dict(
                    symbol=symbol,
                    color=marker_color,
                ),
                mode="lines+markers",
                line=dict(
                    color=line_color,
                    dash=line_style,
                ),
            )
        )
        index += 1

    fig.update_layout(
        polar=dict(
            bgcolor="white",
            gridshape="linear",
            radialaxis=dict(
                range=[0, len(levels) - 1],
                gridcolor="lightgrey",
                angle=90,
                tickangle=90,
            ),
            angularaxis=dict(
                showgrid=False,
                gridcolor="lightgrey",
                tickmode="array",
                linecolor="lightgrey",
                rotation=90,
                direction="clockwise",
            ),
        ),
        plot_bgcolor="rgba(0,0,0,0)",
        showlegend=True,
        legend=dict(
            orientation="h",
            x=0.5,
            y=-0.1,
            xanchor="center",
            yanchor="top",
            traceorder="normal",
            itemwidth=70,
            valign="middle",
        ),
        margin=dict(l=0, r=0, t=50, b=0),
    )

    return fig


def text_wrap(text, max_line_length=20):
    if isinstance(text, list):
        text_wrapped = ["<br>".join(textwrap.wrap(label, width=max_line_length)) for label in text]
    elif isinstance(text, str):
        text_wrapped = "<br>".join(textwrap.wrap(text, width=max_line_length))
    else:
        return None
    return text_wrapped


def convert_graph_to_base64(fig, export_format="png"):
    image_bytes = pio.to_image(fig, format=export_format, scale=2)
    return base64.b64encode(image_bytes).decode("utf-8")


def convert_docx_to_pdf(docx_path: str) -> str:
    docx_path = Path(docx_path).resolve()
    output_dir = docx_path.parent
    pdf_path = output_dir / (docx_path.stem + ".pdf")

    profile_dir = Path(output_dir / f"lo-profile-{uuid.uuid4()}")
    profile_dir.mkdir(parents=True, exist_ok=True)

    try:
        subprocess.run(
            [
                "soffice",
                f"-env:UserInstallation=file://{profile_dir}",
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--nodefault",
                "--norestore",
                "--nofirststartwizard",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=120,
        )

        if not pdf_path.exists():
            raise RuntimeError(f"PDF not created: {pdf_path}")

    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)

    return pdf_path


def merge_subdoc_into_placeholder(main_docx_path, subdoc_path, placeholder, output_path):
    main = Document(main_docx_path)
    sub = Document(subdoc_path)

    _copy_styles(sub, main)

    paragraphs_to_replace = [para for para in main.paragraphs if para.text.strip() == placeholder]

    for para in paragraphs_to_replace:
        parent = para._element.getparent()
        idx = list(parent).index(para._element)
        parent.remove(para._element)

        for j, block in enumerate(sub.element.body):
            if block.tag == qn("w:sectPr"):
                continue
            new_block = copy.deepcopy(block)
            parent.insert(idx + j, new_block)

    main.save(output_path)


def _copy_styles(source_doc, target_doc):
    source_styles = {s.element.get(qn("w:styleId")): s.element for s in source_doc.styles}
    target_style_ids = {s.element.get(qn("w:styleId")) for s in target_doc.styles}

    for style_id, style_elem in source_styles.items():
        if style_id and style_id not in target_style_ids:
            target_doc.styles.element.append(copy.deepcopy(style_elem))


def rgb_to_hex(rgb):
    r, g, b = pc.unlabel_rgb(str(rgb))
    return f"#{int(r):02x}{int(g):02x}{int(b):02x}"


def lighten_color(rgb_base: str, factor: float = 0.5) -> str:
    rgb_lighter = pc.find_intermediate_color(rgb_base, "(255, 255, 255)", factor, colortype="rgb")
    return rgb_to_hex(rgb_lighter)


def get_gradient_color(value, so_color_palette):

    palette = [(float(v), c) for v, c in so_color_palette]

    min_v = palette[0][0]
    max_v = palette[-1][0]

    normalized = (float(value) - min_v) / (max_v - min_v)

    colorscale = [[(v - min_v) / (max_v - min_v), c] for v, c in palette]

    rgb = pc.sample_colorscale(colorscale, [normalized])[0]

    return rgb_to_hex(rgb)


def get_chart_color_palette(colors):
    regulator_palette = [str(f"rgb{pc.hex_to_rgb(color[0])}") for color in colors]
    default_palette = pc.qualitative.Set1 + pc.qualitative.Set2
    return regulator_palette + default_palette


def _get_page_content_width_dxa(doc: Document) -> int:
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin

    content_width_emu = page_width - left_margin - right_margin
    return int(content_width_emu * 1440 / 914400)


def _get_table_total_width(table, doc) -> int:
    DEFAULT_WIDTH_DXA = _get_page_content_width_dxa(doc)

    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        _set_table_total_width(table, DEFAULT_WIDTH_DXA)
        return DEFAULT_WIDTH_DXA

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        _set_table_total_width(table, DEFAULT_WIDTH_DXA)
        return DEFAULT_WIDTH_DXA

    w_type = tblW.get(qn("w:type"))
    w_val = tblW.get(qn("w:w"))

    if w_type == "dxa" and w_val:
        return int(w_val)

    if w_type == "pct" and w_val:
        ratio = int(w_val) / 5000
        width = int(DEFAULT_WIDTH_DXA * ratio)
        _set_table_total_width(table, width)
        return width

    if w_type in ("auto", "nil") or w_val == "0":
        _set_table_total_width(table, DEFAULT_WIDTH_DXA)
        return DEFAULT_WIDTH_DXA

    _set_table_total_width(table, DEFAULT_WIDTH_DXA)
    return DEFAULT_WIDTH_DXA


def _set_table_total_width(table, width_dxa: int):
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)

    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")


def redistribute_column_widths_proportional(table, proportions, doc, table_widht_dxa=None):
    if table_widht_dxa is not None:
        total_width = table_widht_dxa
        _set_table_total_width(table, total_width)
    else:
        total_width = _get_table_total_width(table, doc)

    if total_width is None:
        return

    first_row_tcs = table.rows[0]._tr.findall(qn("w:tc"))
    num_cols = len(first_row_tcs)

    if proportions is None:
        proportions = [1 / num_cols] * num_cols

    total = sum(proportions)
    proportions = [p / total for p in proportions]
    col_widths = [int(total_width * p) for p in proportions]
    col_widths[-1] = total_width - sum(col_widths[:-1])

    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    for row in table.rows:
        tcs = row._tr.findall(qn("w:tc"))
        for i, tc in enumerate(tcs):
            if i >= len(col_widths):
                continue
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(col_widths[i]))
            tcW.set(qn("w:type"), "dxa")

    tblGrid = table._element.find(qn("w:tblGrid"))
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn("w:gridCol"))
        for i, gridCol in enumerate(gridCols):
            if i < len(col_widths):
                gridCol.set(qn("w:w"), str(col_widths[i]))


def fix_outer_column_borders(table_element, color: str = "FFFFFF"):
    for row in table_element.findall(qn("w:tr")):
        tcs = row.findall(qn("w:tc"))
        if not tcs:
            continue

        cells_sides = [
            (tcs[0], "left"),
            (tcs[-1], "right"),
        ]

        for tc, side in cells_sides:
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)

            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)

            existing = tcBorders.find(qn(f"w:{side}"))
            if existing is not None:
                tcBorders.remove(existing)

            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color)
            tcBorders.append(border)


def is_soffice_running(pipe_name="update_toc"):
    """Check if soffice is running with the correct pipe"""
    for proc in psutil.process_iter(["name", "cmdline"]):
        try:
            if proc.info["name"] and "soffice" in proc.info["name"]:
                cmdline = proc.info["cmdline"] or []
                if any(pipe_name in arg for arg in cmdline):
                    return True
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    return False


def ensure_soffice_running(pipe_name: str = "update_toc"):
    if not is_soffice_running(pipe_name):
        subprocess.Popen(
            [
                "soffice",
                "--headless",
                f"--accept=pipe,name={pipe_name};urp;",
                "--norestore",
                "--nologo",
                "--invisible",
            ],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )


# replace ToC number by the one of array otc_number
def replace_toc_page_numbers(doc_v1: Document, toc_number) -> None:
    index_number = dict()
    title_number = dict()
    for entry in toc_number:
        page = entry["page"]
        if entry["index"] is not None:
            index_number[entry["index"]] = page
        if entry["name"] is not None:
            title_number[entry["name"]] = page

    body = doc_v1.element.body
    paragraphs = list(body.iter(qn("w:p")))
    in_toc = False

    for para in paragraphs:
        title = None
        pStyle = para.find(".//" + qn("w:pStyle"))
        style_val = pStyle.get(qn("w:val"), "") if pStyle is not None else ""

        if not in_toc:
            for instr in para.iter(qn("w:instrText")):
                if "TOC" in (instr.text or ""):
                    in_toc = True
                    break

        if not in_toc:
            continue

        # Detect end of Toc paragraphe NO_STYLE with fldChar[end] of TOC global field
        if not style_val.startswith("TOC") and not style_val.startswith("toc"):
            has_toc_instr = any("TOC" in (instr.text or "") for instr in para.iter(qn("w:instrText")))
            if not has_toc_instr:
                # check if it's the fldChar[end] of global ToC
                fld_chars = list(para.iter(qn("w:fldChar")))
                if any(fc.get(qn("w:fldCharType")) == "end" for fc in fld_chars):
                    break
                # we continue
                if not any(para.iter(qn("w:hyperlink"))):
                    continue
        title = extract_title_from_para(para)
        index = extract_index_from_para(para)
        # in each ToC, empty the w:t after fldChar[separate] of PAGEREF
        if (title and title in title_number) or (index and index in index_number):
            for hyperlink in para.iter(qn("w:hyperlink")):
                in_pageref_value = False
                for r in hyperlink.iter(qn("w:r")):
                    fld_char = r.find(qn("w:fldChar"))
                    if fld_char is not None:
                        fld_type = fld_char.get(qn("w:fldCharType"))
                        if fld_type == "separate":
                            in_pageref_value = True
                        elif fld_type == "end":
                            in_pageref_value = False

                    if in_pageref_value:
                        t = r.find(qn("w:t"))
                        if t is not None:
                            if index in index_number:
                                t.text = index_number[index]
                            elif title in title_number:
                                t.text = title_number[title]


# extract the index of a docx paragraph
def extract_index_from_para(para):
    texts = [t.text.strip() for t in para.iter(qn("w:t")) if t.text and t.text.strip()]

    if len(texts) < 2:
        return None

    # remove last (page number)
    texts_no_page = texts[:-1]

    # index = premier élément numérique du TOC
    for t in texts_no_page:
        if re.match(r"^\d+(\.\d+)*$", t):
            return t

    return None


# extract the title of a docx paragraph
def extract_title_from_para(para):
    texts = [t.text.strip() for t in para.iter(qn("w:t")) if t.text and t.text.strip()]

    if len(texts) < 2:
        return None

    # remove the last (page number)
    texts_no_page = texts[:-1]

    # remove the begning (ex : 1.2.3)
    clean_texts = [t for t in texts_no_page if not re.match(r"^\d+(\.\d+)*$", t)]

    if not clean_texts:
        return None

    # real title = last one
    texte = " ".join(clean_texts)
    return " ".join(texte.split())


# update page number with libre office
def get_updated_toc(file_path: str, pipe_name: str = "update_toc"):
    """
    Update table of content in docx
    """

    cmd = [
        "/usr/bin/python3",
        "reporting/scripts/update_toc_docx.py",
        file_path,
        pipe_name,
    ]

    proc = subprocess.run(
        cmd,
        check=True,
        timeout=120,
        capture_output=True,
        text=True,
    )

    return json.loads(proc.stdout)
