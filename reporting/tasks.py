import base64
import datetime
import json
import os
import shutil
import uuid
import zipfile
from decimal import Decimal
from io import BytesIO
from pathlib import Path

from celery import shared_task
from celery.utils.log import get_task_logger
from django.conf import settings
from django.contrib.auth import get_user_model
from django.utils import formats
from django.utils.translation import activate
from docx import Document
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage, RichTextParagraph

from .globals import CELERY_TASK_STATUS, TRANSLATIONS_CONTEXT
from .helpers import (
    convert_docx_to_pdf,
    create_entry_log,
    ensure_soffice_running,
    fix_outer_column_borders,
    get_charts,
    get_risk_data,
    get_so_data,
    get_updated_toc,
    merge_subdoc_into_placeholder,
    redistribute_column_widths_proportional,
    replace_toc_page_numbers,
)
from .models import Configuration, GeneratedReport, Project, Template

logger = get_task_logger(__name__)


@shared_task(bind=True, ignore_result=True)
def generate_data(self, cleaned_data):
    def custom_default(obj):
        if isinstance(obj, Decimal):
            return float(obj)
        if isinstance(obj, uuid.UUID):
            return str(obj)
        if isinstance(obj, (datetime.datetime, datetime.date)):
            return obj.isoformat()
        return str(obj)

    project_id = cleaned_data["project_id"]
    project = Project.objects.get(id=project_id)
    if project.task_status == "ABORT" or project.task_status == "FAIL":
        return "Aborted"
    run_id = str(self.request.root_id)
    base_tmp_dir = Path(settings.PATH_FOR_REPORTING_PDF)
    task_tmp_dir = (
        base_tmp_dir / str(project_id) / f"tmp_files_{project.task_id}" / run_id
    )
    file_path = task_tmp_dir / "data.json"
    task_tmp_dir.mkdir(parents=True, exist_ok=True)
    language = cleaned_data.get("language", "en")
    activate(language)
    report_configuration_id = cleaned_data["report_configuration_id"]
    colors = Configuration.objects.get(pk=report_configuration_id).colors.values_list(
        "color"
    )
    template_id = cleaned_data["template_id"]
    so_data = get_so_data(cleaned_data)
    risk_data = get_risk_data(cleaned_data)
    charts = get_charts(so_data, risk_data, colors)

    data = {
        "company": cleaned_data["company"]["name"],
        "years": cleaned_data["years"],
        "reference_year": cleaned_data["reference_year"],
        "sector": cleaned_data["sector"]["name"],
        "threshold_for_high_risk": cleaned_data["threshold_for_high_risk"],
        "top_ranking": cleaned_data["top_ranking"],
        "report_recommendations": cleaned_data["report_recommendations"],
        "charts": charts,
        "so_data": so_data,
        "risk_data": risk_data,
        "company_reporting": cleaned_data["company_reporting"],
        "translations": {k: str(v) for k, v in TRANSLATIONS_CONTEXT.items()},
        "template_id": template_id,
        "project_id": project_id,
    }

    with open(file_path, "w") as f:
        json.dump(data, f, default=custom_default)

    return "Data generated successfully"


@shared_task(
    bind=True,
    ignore_result=True,
    autoretry_for=(Exception,),
    max_retries=3,
    retry_backoff=True,
    retry_backoff_max=120,
    retry_jitter=True,
)
def generate_docx_task(self, project_id):
    project = Project.objects.get(id=project_id)
    if project.task_status == "ABORT" or project.task_status == "FAIL":
        return "Aborted"
    run_id = str(self.request.root_id)
    base_tmp_dir = Path(settings.PATH_FOR_REPORTING_PDF)
    task_tmp_dir = (
        base_tmp_dir / str(project_id) / f"tmp_files_{project.task_id}" / run_id
    )
    file_path = Path(task_tmp_dir / "data.json")

    if not file_path.exists():
        return "Data file not found for docx generation"

    with open(file_path) as f:
        data = json.load(f)

    if not data:
        return "No data found for docx generation"

    if Project.objects.get(id=project_id).task_status == "ABORT":
        return "Aborted"
    subdocs_templates_dir = Path(
        os.path.join(settings.BASE_DIR, "reporting", "subdocs_templates")
    )
    success = False
    try:
        if not task_tmp_dir.exists():
            logger.warning(
                "task_tmp_dir has disappeared before processing begins : %s",
                task_tmp_dir,
            )
            return
        template_id = data["template_id"]
        template_file = Template.objects.get(pk=template_id).template_file
        template_path = BytesIO(bytes(template_file))
        rendered_subs_docs = {}
        nb_years = len(data["years"])
        document_charts = {
            "chart_average_risk_level": {
                "width": Mm(140),
            },
            "chart_high_risk_rate": {
                "width": Mm(140),
            },
            "chart_average_high_risk_level": {
                "width": Mm(140),
            },
            "chart_evolution_highest_risks": {
                "width": Mm(140),
            },
        }
        document_tables = {
            "table_of_evolution_security_objectives": {
                "context": {
                    "table": data["so_data"]["company_so_by_year"],
                },
                "column_proportions": [0.4] + [0.1] * nb_years + [0.15],
            },
            "table_of_evolution_security_objectives_by_domain": {
                "context": {
                    "table": data["so_data"]["company_so_by_domain"],
                },
                "column_proportions": [0.4] + [0.1] * nb_years + [0.15] * 2,
            },
            "table_of_highest_security_objectives_in_the_sector": {
                "context": {
                    "table": data["so_data"]["sector_so_by_year_desc"][
                        str(data["reference_year"])
                    ],
                },
                "column_proportions": [0.05] + [0.65] + [0.15] * 2,
            },
            "table_of_lowest_security_objectives_in_the_sector": {
                "context": {
                    "table": data["so_data"]["sector_so_by_year_asc"][
                        str(data["reference_year"])
                    ],
                },
                "column_proportions": [0.05] + [0.65] + [0.15] * 2,
            },
            "table_of_evolution_of_the_weakest_security_objectives": {
                "context": {
                    "table": data["so_data"]["company_so_by_priority"],
                },
                "column_proportions": [0.4] + [0.15] * nb_years + [0.15],
            },
            "table_of_security_objectives_by_maturity_level": {
                "context": {
                    "table": data["so_data"]["company_so_by_level"],
                },
                "column_proportions": [0.25]
                * len(data["so_data"]["company_so_by_level"]["headers"]),
            },
            "maturity_level_legend": {
                "context": {
                    "maturity_levels": data["so_data"]["maturity_levels"],
                },
            },
            "table_of_evolution_of_the_highest_risks": {
                "context": {
                    "table": data["risk_data"]["data_risks_top_ranking"],
                },
                "column_proportions": [0.1, 0.25, 0.25, 0.3] + [0.1] * nb_years,
            },
            "table_of_treatment_of_the_highest_risks": {
                "context": {
                    "table": data["risk_data"]["data_risks_top_ranking"],
                },
                "column_proportions": [0.08, 0.18, 0.18, 0.25, 0.13, 0.17, 0.13],
            },
            "table_of_risk_summary": {
                "context": {
                    "table": data["risk_data"]["risks_stats_by_year"],
                },
                "column_proportions": [0.7] + [0.15] * nb_years,
                "table_width_dxa": 7380,  # 13cm
            },
            "table_of_top_threats_by_occurrence": {
                "context": {
                    "table": data["risk_data"]["top_threats"],
                },
                "column_proportions": [0.1, 0.9],
                "table_width_dxa": 8504,  # 15cm
            },
            "table_of_top_vulnerabilities_by_occurrence": {
                "context": {
                    "table": data["risk_data"]["top_vulnerabilities"],
                },
                "column_proportions": [0.1, 0.9],
                "table_width_dxa": 8504,  # 15cm
            },
            "table_of_recommendations": {
                "context": {
                    "table": data["risk_data"]["recommendations_evolution"],
                },
                "column_proportions": [0.7, 0.15, 0.15],
            },
        }

        main_doc_template = DocxTemplate(template_path)
        main_doc = Document(template_path)

        report_recommendations = RichTextParagraph()
        for rec in data["report_recommendations"]:
            report_recommendations.add(rec, parastyle="CircleBullet")

        context = {
            "operator_name": data["company"],
            "sector": data["sector"],
            "year": data["reference_year"],
            "threshold_for_high_risk": data["threshold_for_high_risk"],
            "top_ranking": data["top_ranking"],
            "report_recommendations": report_recommendations,
            "report_observations": data["company_reporting"]["comment"],
            "publication_date": formats.date_format(
                datetime.date.today(), format="d F Y"
            ),
        }
        for chart_name, chart_data in data["charts"].items():
            chart_bytes = BytesIO(base64.b64decode(chart_data))
            chart_with = document_charts.get(chart_name, {}).get("width", Mm(170))
            context[chart_name] = InlineImage(
                main_doc_template, chart_bytes, width=chart_with
            )

        for table_name, table_info in document_tables.items():
            sub_template_path = subdocs_templates_dir / f"{table_name}_template.docx"
            sub_rendered_path = task_tmp_dir / f"{table_name}_rendered.docx"
            context[table_name] = str(table_name)
            table_info["context"].update(
                {
                    "translations": data["translations"],
                    "year": data["reference_year"],
                    "years": data["years"],
                }
            )
            sub_doc_template = DocxTemplate(sub_template_path)
            sub_doc_template.render(table_info["context"])
            sub_doc_template.save(sub_rendered_path)
            sub_doc = Document(sub_rendered_path)
            for table in sub_doc.tables:
                table_width_dxa = None
                if "table_width_dxa" in table_info:
                    # 1dxa = 1cm * 28.346 * 20
                    table_width_dxa = table_info["table_width_dxa"]
                if "column_proportions" in table_info:
                    redistribute_column_widths_proportional(
                        table,
                        table_info["column_proportions"],
                        main_doc,
                        table_width_dxa,
                    )
                fix_outer_column_borders(table._element)
            sub_doc.save(sub_rendered_path)
            rendered_subs_docs[table_name] = sub_rendered_path

        main_docx_path = Path(task_tmp_dir / "main_doc.docx")
        main_doc_template.render(context)
        main_doc_template.save(main_docx_path)
        current_doc = main_docx_path
        tmp_output_path = task_tmp_dir / "tmp_doc.docx"

        for placeholder, sub_rendered_path in rendered_subs_docs.items():
            sub_rendered_path = Path(sub_rendered_path)
            try:
                merge_subdoc_into_placeholder(
                    main_docx_path=current_doc,
                    subdoc_path=sub_rendered_path,
                    placeholder=placeholder,
                    output_path=tmp_output_path,
                )
                current_doc = tmp_output_path
            finally:
                if sub_rendered_path.exists():
                    sub_rendered_path.unlink(missing_ok=True)

        pipe_name = f"update_toc_{os.getpid()}"
        ensure_soffice_running(pipe_name)
        # get updated ToC
        updated_toc = get_updated_toc(str(current_doc), pipe_name)
        doc = Document(str(current_doc))
        replace_toc_page_numbers(doc, updated_toc)
        doc.save(str(current_doc))
        main_docx_path.unlink(missing_ok=True)
        success = True
        return "Docx generated successfully"

    except Exception as exc:
        logger.exception(
            "Error in generate_docx_task for the project %s : %s", project_id, exc
        )
        raise
    finally:
        if success and task_tmp_dir.exists():
            file_path.unlink(missing_ok=True)
        if not success and task_tmp_dir.exists():
            shutil.rmtree(task_tmp_dir, ignore_errors=True)
            logger.info("Temporary file cleanup following a failure : %s", task_tmp_dir)


@shared_task(
    bind=True,
    ignore_result=True,
    autoretry_for=(Exception,),
    max_retries=3,
    retry_backoff=True,
    retry_backoff_max=120,
    retry_jitter=True,
)
def generate_pdf_task(self, project_id):
    project = Project.objects.get(id=project_id)
    if project.task_status == "ABORT" or project.task_status == "FAIL":
        return "Aborted"
    run_id = str(self.request.root_id)
    base_tmp_dir = Path(settings.PATH_FOR_REPORTING_PDF)
    task_tmp_dir = (
        base_tmp_dir / str(project_id) / f"tmp_files_{project.task_id}" / run_id
    )
    docx_path = Path(task_tmp_dir / "tmp_doc.docx")

    if not docx_path.exists():
        return "Docx file not found for PDF conversion"
    try:
        convert_docx_to_pdf(str(docx_path))
        docx_path.unlink(missing_ok=True)
        return "PDF generated successfully"
    except Exception as exc:
        logger.exception("PDF conversion failed for project %s : %s", project_id, exc)
        raise


@shared_task(bind=True, ignore_result=True)
def save_file_task(self, project_id, user_id, filename, is_multiple_files):
    project = Project.objects.get(id=project_id)
    if project.task_status == "ABORT" or project.task_status == "FAIL":
        return "Aborted"
    run_id = str(self.request.root_id)
    base_tmp_dir = Path(settings.PATH_FOR_REPORTING_PDF)
    task_tmp_dir = Path(
        base_tmp_dir / str(project_id) / f"tmp_files_{project.task_id}" / run_id
    )
    temp_file_path = next(task_tmp_dir.glob("tmp_doc.*"), None)
    if temp_file_path is None:
        logger.warning("Temporary file not found in %s", task_tmp_dir)
        return

    file_uuid = uuid.uuid4()
    User = get_user_model()
    user = User.objects.get(id=user_id)
    output_dir = Path(os.path.join(settings.PATH_FOR_REPORTING_PDF, str(project_id)))
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, str(file_uuid))

    if is_multiple_files:
        file_path = os.path.join(temp_file_path.parent, filename)
        shutil.move(temp_file_path, file_path)

    else:
        shutil.move(temp_file_path, file_path)
        GeneratedReport.objects.update_or_create(
            project=project,
            defaults={"file_uuid": file_uuid, "filename": filename},
        )
        create_entry_log(user, project, "GENERATE REPORT")
        project.task_status = "DONE"
        project.save()

    return "File saved successfully"


@shared_task(ignore_result=True)
def zip_files_task(user_id, project_id, error_messages):
    project = Project.objects.get(id=project_id)
    if project.task_status == "ABORT" or project.task_status == "FAIL":
        return "Aborted"
    User = get_user_model()
    user = User.objects.get(id=user_id)
    file_uuid = uuid.uuid4()
    output_dir = Path(os.path.join(settings.PATH_FOR_REPORTING_PDF, str(project_id)))
    output_dir.mkdir(parents=True, exist_ok=True)

    zip_filename = f"reports_{project.name}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    zip_path = os.path.join(output_dir, str(file_uuid))

    base_tmp_dir = Path(settings.PATH_FOR_REPORTING_PDF)
    task_tmp_dir = Path(base_tmp_dir / str(project_id) / f"tmp_files_{project.task_id}")

    with zipfile.ZipFile(zip_path, "w") as zipf:
        for run_path in task_tmp_dir.iterdir():
            if not run_path.is_dir():
                continue
            file_path = next(
                (f for f in run_path.iterdir() if f.suffix in (".pdf", ".docx")), None
            )
            if file_path and file_path.exists():
                arcname = os.path.basename(file_path)
                zipf.write(file_path, arcname=arcname)
            else:
                logger.warning("File missing from the zip file : %s", file_path)

        if error_messages:
            error_log = "\n".join(error_messages)
            zipf.writestr("error_log.txt", error_log)

    GeneratedReport.objects.update_or_create(
        project=project,
        defaults={"file_uuid": file_uuid, "filename": zip_filename},
    )
    create_entry_log(user, project, "GENERATE REPORT")
    project.task_status = "DONE"
    project.save()

    return "Files zipped and saved successfully"


@shared_task(ignore_result=True)
def cleanup_files(project_id, all_files=False):
    base_tmp_dir = Path(settings.PATH_FOR_REPORTING_PDF)
    task_tmp_dir = base_tmp_dir / str(project_id)

    if not task_tmp_dir.exists():
        return "Temporary directory does not exist, no files to clean up"

    if all_files:
        shutil.rmtree(task_tmp_dir, ignore_errors=True)
        return "Temporary files cleaned up successfully"

    items = list(task_tmp_dir.iterdir())

    if len(items) <= 1:
        return "No files to clean up"

    items.sort(key=lambda x: x.stat().st_mtime)
    latest_item = items[-1]

    if latest_item.is_dir():
        shutil.rmtree(latest_item, ignore_errors=True)
        return "Temporary files cleaned up successfully"

    for item in items:
        if item != latest_item:
            if item.is_dir():
                shutil.rmtree(item, ignore_errors=True)
            else:
                item.unlink(missing_ok=True)

    return "Temporary files cleaned up successfully"


@shared_task(ignore_result=True)
def on_chord_error(request, exc, traceback, project_id):
    logger.error("Generation failed for the project %s : %s", project_id, exc)
    Project.objects.filter(id=project_id).update(task_status=CELERY_TASK_STATUS[0][0])
    cleanup_files.delay(project_id)
